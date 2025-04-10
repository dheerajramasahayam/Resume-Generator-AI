from docx import Document
from docx.shared import Pt, Inches, RGBColor # Import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement # For borders
from .common import is_heading # Import common helper

def _add_paragraph_border(paragraph):
    """Helper to add a bottom border to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    p_pr.append(p_borders)
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6') # Border size (eighths of a point)
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    p_borders.append(bottom_bdr)

def generate_simple_docx(document, cleaned_text):
    """Applies simple DOCX formatting."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    # Set narrow margins for simple
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)


    for paragraph_text in cleaned_text.split('\n'):
        line_stripped = paragraph_text.strip()
        if is_heading(line_stripped):
             # Simple heading: slightly larger, bold
             paragraph = document.add_paragraph()
             run = paragraph.add_run(line_stripped)
             run.font.name = 'Calibri'
             run.font.size = Pt(12)
             run.bold = True
             paragraph_format = paragraph.paragraph_format
             paragraph_format.space_before = Pt(8)
             paragraph_format.space_after = Pt(4)
        elif line_stripped:
            paragraph = document.add_paragraph(line_stripped)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(4) # Less spacing for simple
        elif len(document.paragraphs) > 0 and document.paragraphs[-1].text.strip():
             # Add a single blank paragraph for spacing if previous wasn't blank
             document.add_paragraph()

def generate_classic_docx(document, cleaned_text):
    """Applies classic DOCX formatting."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11) # Standard body size
    # Set standard margins for classic
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Define heading style if not already defined (basic check)
    try:
        heading_style = document.styles['ResumeHeadingClassic'] # Unique name
    except KeyError:
        heading_style = document.styles.add_style('ResumeHeadingClassic', 1)
        heading_font = heading_style.font
        heading_font.name = 'Times New Roman'
        heading_font.size = Pt(13)
        heading_font.bold = True
        # Define a style for bullet points if needed, or use default list style
        try:
            bullet_style = document.styles['List Bullet']
        except KeyError:
             # If 'List Bullet' doesn't exist, create or use a fallback
             try:
                 bullet_style = document.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
                 bullet_style.base_style = document.styles['Normal']
                 # Add indentation if needed:
                 # bullet_style.paragraph_format.left_indent = Inches(0.25)
             except: # Fallback if style creation fails
                 bullet_style = document.styles['Normal']


    # Process lines based on markdown markers
    for paragraph_text in cleaned_text.split('\n'):
        line_stripped = paragraph_text.strip()

        if line_stripped.startswith('### '):
            heading_text = line_stripped[4:].strip()
            paragraph = document.add_paragraph(heading_text.upper(), style='ResumeHeadingClassic')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)
            _add_paragraph_border(paragraph)
        elif line_stripped.startswith('* '):
            item_text = line_stripped[2:].strip()
            # Applying 'List Bullet' style handles indentation/bullet
            paragraph = document.add_paragraph(item_text, style='List Bullet')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(2) # Tighter spacing for bullets
        elif '**' in line_stripped: # Basic check for bold text (job titles etc.)
             paragraph = document.add_paragraph()
             parts = re.split(r'(\*\*.*?\*\*)', line_stripped) # Split by bold markers
             for part in parts:
                 if part.startswith('**') and part.endswith('**'):
                     run = paragraph.add_run(part[2:-2])
                     run.bold = True
                 elif part: # Add non-bold parts
                     paragraph.add_run(part)
             paragraph_format = paragraph.paragraph_format
             paragraph_format.space_after = Pt(6) # Standard spacing
        elif line_stripped: # Regular paragraph
            paragraph = document.add_paragraph(line_stripped)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(6)
        elif len(document.paragraphs) > 0 and document.paragraphs[-1].text.strip():
             last_para_format = document.paragraphs[-1].paragraph_format
             if last_para_format.space_after is None or last_para_format.space_after < Pt(10):
                 document.add_paragraph()

def generate_modern_docx(document, cleaned_text):
    """Applies modern DOCX formatting (example)."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10) # Smaller body font for modern
     # Set slightly wider side margins for modern
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Define heading style if not already defined
    try:
        heading_style = document.styles['ResumeHeadingModern']
    except KeyError:
        heading_style = document.styles.add_style('ResumeHeadingModern', 1)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(12)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x2E, 0x86, 0xC1) # Add a blue color

    for paragraph_text in cleaned_text.split('\n'):
        line_stripped = paragraph_text.strip()
        if is_heading(line_stripped):
            paragraph = document.add_paragraph(line_stripped.upper(), style='ResumeHeadingModern') # UPPERCASE MODERN HEADINGS
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(12) # More space before modern heading
            paragraph_format.space_after = Pt(3)  # Tight space after modern heading
            # No border for modern heading style
        elif line_stripped:
            paragraph = document.add_paragraph(line_stripped)
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_after = Pt(4) # Less spacing for modern body
        elif len(document.paragraphs) > 0 and document.paragraphs[-1].text.strip():
             document.add_paragraph() # Simpler spacing for modern

# Add more template functions as needed
